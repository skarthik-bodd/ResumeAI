from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from .types import Document


SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}


class DocumentLoadError(ValueError):
    """Raised when input files cannot be read."""


def discover_files(inputs: list[str | Path]) -> list[Path]:
    files: list[Path] = []
    seen: set[Path] = set()

    for item in inputs:
        path = Path(item).expanduser().resolve()
        if not path.exists():
            raise DocumentLoadError(f"Input path does not exist: {path}")

        candidates: list[Path]
        if path.is_dir():
            candidates = [p for p in path.rglob("*") if p.is_file()]
        else:
            candidates = [path]

        for candidate in candidates:
            ext = candidate.suffix.lower()
            if ext not in SUPPORTED_EXTENSIONS:
                continue
            if candidate in seen:
                continue
            seen.add(candidate)
            files.append(candidate)

    if not files:
        raise DocumentLoadError(
            "No supported files found. Supported extensions: .md, .txt, .pdf, .docx"
        )

    files.sort()
    return files


def read_file_text(path: str | Path) -> str:
    target = Path(path)
    ext = target.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise DocumentLoadError(f"Unsupported file type: {target}")

    if ext in {".md", ".txt"}:
        return target.read_text(encoding="utf-8", errors="ignore")

    if ext == ".pdf":
        return _read_pdf(target)

    if ext == ".docx":
        return _read_docx(target)

    raise DocumentLoadError(f"Unsupported file type: {target}")


def load_documents(inputs: list[str | Path]) -> list[Document]:
    documents: list[Document] = []
    for path in discover_files(inputs):
        text = read_file_text(path).strip()
        if not text:
            continue
        documents.append(Document(source=str(path), text=text))

    if not documents:
        raise DocumentLoadError("No readable content found in supported files.")

    return documents


def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


def _read_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)
